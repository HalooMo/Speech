"""
Генерация 4 документов из people.json по шаблонам Pattern/.
Результат: {work_dir}/{ДД.ММ.ГГГГ}/*.docx|xlsx|doc
"""

import re
import shutil
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document

import storage

APP_DIR = Path(__file__).resolve().parent
PATTERN_DIR = APP_DIR / "Pattern"

TEMPLATE_FULL_XLSX = PATTERN_DIR / "Строевая записка 02.06.2026.xlsx"
TEMPLATE_NARROW_XLSX = PATTERN_DIR / "Строевая записка 19.05.2026 1125 Ц(РХБЗ).xlsx"
TEMPLATE_STROEVAYA_DOCX = PATTERN_DIR / "Строевая записка 19.05.docx"
TEMPLATE_SHDS_DOC = PATTERN_DIR / "ШДС СО 30 чел (ОРФОШ).doc"

# Строки таблицы 1 в docx «Строевая» (0-based)
DOCX_REASON_ROW = {
    "наряд": 1,
    "командировка": 2,
    "отпуск": 4,
    "40 РЦПС": 11,
    "больничный": 7,
    "госпиталь": 8,
    "прочие причины": 9,
}
DOCX_ON_DUTY_ROWS = list(range(12, 18))  # «на лицо» — 6 строк

# Пары столбцов в/с, г/п в таблице 0 (0-based col index)
DOCX_COUNT_COL = {
    "наряд": 12,
    "командировка": 14,
    "отпуск": 16,
    "прочие причины": 20,
    "госпиталь": 22,
    "больничный": 24,
}

# Excel: те же смещения от колонки M (13) как в шаблоне R16
XLSX_COUNT_COL = {
    "наряд": 14,      # N
    "командировка": 16,  # P
    "отпуск": 18,     # R
    "прочие причины": 20,  # T
    "госпиталь": 22,  # V
    "больничный": 24,  # X
}

SHDS_RHBZ_ROWS = (13, 14, 15, 16, 17, 18)  # 6 должностей группы РХБЗ
SHDS_DEFAULT_POSITIONS = [
    "Командир группы",
    "Химик-дозиметрист",
    "Водитель-химик",
    "Химик-дозиметрист",
    "Химик-дозиметрист",
    "Химик-дозиметрист",
]


def today_str() -> str:
    return datetime.now().strftime("%d.%m.%Y")


def output_dir() -> Path:
    work = storage.get_work_dir()
    if not work:
        raise RuntimeError(storage.require_work_dir()[1])
    folder = work / today_str()
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _is_officer(rank: str) -> bool:
    r = rank.lower()
    keys = ("майор", "капитан", "полковник", "подполковник", "лейтенант", "л-т", "м-р", "к-н", "п-к", "полковник", "генерал")
    return any(k in r for k in keys)


def _is_prapor(rank: str) -> bool:
    r = rank.lower()
    return "прапор" in r or "пр-к" in r or "пр." in r


# Сокращения званий как в строевой (ряд. Лужбин, м-р Байдола)
RANK_SHORT = {
    "рядовой": "ряд.",
    "ефрейтор": "ефр.",
    "младший сержант": "мл. с-т",
    "сержант": "с-т",
    "старшина": "ст. с-т",
    "прапорщик": "пр-к",
    "старший прапорщик": "ст. пр-к",
    "мичман": "мичм.",
    "лейтенант": "л-т",
    "старший лейтенант": "ст. л-т",
    "капитан": "к-н",
    "майор": "м-р",
    "подполковник": "п/п-к",
    "полковник": "п-к",
    "генерал-майор": "ген.-майор",
}


def _short_rank(rank: str) -> str:
    r = rank.strip().lower().replace("ё", "е")
    if not r:
        return ""
    for full, short in RANK_SHORT.items():
        if full in r or r == full:
            return short
    # уже сокращение (м-р, ряд., …)
    if len(r) <= 6:
        return rank.strip()
    return rank.strip()


def format_person_line_stroevaya(p: dict) -> str:
    """Формат шаблона: «ряд. Лужбин», «м-р Байдола»."""
    fio = p.get("fio", "").strip()
    parts = fio.split()
    surname = parts[0] if parts else fio
    rank = _short_rank(p.get("rank", ""))
    return f"{rank} {surname}".strip()


def format_person_line(p: dict) -> str:
    """Для docx: звание + фамилия и инициалы."""
    rank = p.get("rank", "").strip()
    fio = p.get("fio", "").strip()
    parts = fio.split()
    if len(parts) >= 3:
        short = f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    elif len(parts) == 2:
        short = f"{parts[0]} {parts[1][0]}."
    else:
        short = fio
    return f"{rank} {short}".strip() if rank else short


def group_by_status(people: list[dict]) -> dict[str, list[dict]]:
    groups = {r: [] for r in storage.REASONS}
    groups[""] = []
    for p in people:
        st = p.get("status", "")
        if st in groups:
            groups[st].append(p)
        else:
            groups[""].append(p)
    return groups


def _column_for_person(p: dict) -> int:
    """Столбец табл.1: 1 офицеры, 2 прапорщики, 3 контрактники."""
    rank = p.get("rank", "")
    if _is_officer(rank):
        return 1
    if _is_prapor(rank):
        return 2
    return 3


def _join_names(people: list[dict]) -> str:
    if not people:
        return ""
    return ", ".join(format_person_line(p) for p in people)


def _split_vertical(people: list[dict], parts: int) -> list[str]:
    """Разбить список на parts ячеек (вертикально «на лицо»)."""
    if not people:
        return [""] * parts
    chunks = [[] for _ in range(parts)]
    for i, p in enumerate(people):
        chunks[i % parts].append(p)
    return [_join_names(c) for c in chunks]


def fill_stroevaya_docx(dst: Path, people: list[dict], date_s: str) -> None:
    doc = Document(TEMPLATE_STROEVAYA_DOCX)
    groups = group_by_status(people)

    # Заголовок с датой
    for para in doc.paragraphs[:6]:
        if "2026" in para.text or "апрел" in para.text.lower():
            para.text = re.sub(
                r"на\s+\d+\s+\S+\s+\d{4}",
                f"на {date_s}",
                para.text,
                count=1,
            )

    t0 = doc.tables[0]
    data_row = 3
    # Счётчики по причинам (в/с; г/п = 0)
    for reason, col in DOCX_COUNT_COL.items():
        n = len(groups.get(reason, []))
        t0.cell(data_row, col).text = str(n)
        t0.cell(data_row, col + 1).text = "0"

    # На лицо
    working = groups[""]
    t0.cell(data_row, 9).text = str(len(working))
    t0.cell(data_row, 10).text = "0"
    t0.cell(data_row, 11).text = str(len(working))

    t1 = doc.tables[1]
    # Очистить ячейки ФИО (колонки 1–4), кроме заголовка
    for ri in range(1, len(t1.rows)):
        for ci in range(1, 5):
            t1.cell(ri, ci).text = ""

    for reason, row_i in DOCX_REASON_ROW.items():
        chunk = groups.get(reason, [])
        by_col = {1: [], 2: [], 3: []}
        for p in chunk:
            by_col[_column_for_person(p)].append(p)
        for ci, lst in by_col.items():
            if lst:
                t1.cell(row_i, ci).text = _join_names(lst)
        total = len(chunk)
        t1.cell(row_i, 5).text = f"{total}/0" if total else "0/0"

    # На лицо — вертикально в строках 12–17
    officers = [p for p in working if _column_for_person(p) == 1]
    prapors = [p for p in working if _column_for_person(p) == 2]
    contract = [p for p in working if _column_for_person(p) == 3]

    off_lines = _split_vertical(officers, len(DOCX_ON_DUTY_ROWS))
    pra_lines = _split_vertical(prapors, len(DOCX_ON_DUTY_ROWS))
    con_lines = _split_vertical(contract, len(DOCX_ON_DUTY_ROWS))

    for i, row_i in enumerate(DOCX_ON_DUTY_ROWS):
        if off_lines[i]:
            t1.cell(row_i, 1).text = off_lines[i]
        if pra_lines[i]:
            t1.cell(row_i, 2).text = pra_lines[i]
        if con_lines[i]:
            t1.cell(row_i, 3).text = con_lines[i]

    doc.save(dst)


def _find_1125_row(ws, col_b: int = 2) -> int | None:
    for r in range(1, ws.max_row + 1):
        v = ws.cell(r, col_b).value
        if v and "1125" in str(v):
            return r
    return None


def _set_xlsx_counts(ws, row: int, groups: dict[str, list[dict]]) -> None:
    for reason, col in XLSX_COUNT_COL.items():
        n = len(groups.get(reason, []))
        ws.cell(row, col).value = n
        ws.cell(row, col + 1).value = 0
    working = groups[""]
    ws.cell(row, 11).value = len(working)  # K на лицо в/с
    ws.cell(row, 12).value = 0


def _build_komandirovka_text(people: list[dict]) -> str:
    """Список командировки 1125 ЦРХБЗ — как в шаблоне, через запятую."""
    parts = [format_person_line_stroevaya(p) for p in people]
    return ", ".join(parts) if parts else ""


def _write_lист2_komandirovka(ws, row_1125: int, kom: list[dict]) -> None:
    """
    Верхняя ячейка B (УТЦ, 46 КЦ из шаблона) не трогаем.
    Командировку 1125 пишем в B на следующей строке.
    """
    kom_row = row_1125 + 1
    if kom:
        ws.cell(kom_row, 2).value = _build_komandirovka_text(kom)
        ws.cell(kom_row, 3).value = len(kom)
    # B[row_1125] и C[row_1125] остаются из шаблона (оформление УТЦ / 46 КЦ)


def fill_full_xlsx(dst: Path, people: list[dict], date_s: str) -> None:
    shutil.copy2(TEMPLATE_FULL_XLSX, dst)
    wb = openpyxl.load_workbook(dst)
    groups = group_by_status(people)

    ws1 = wb["Лист1"]
    row = _find_1125_row(ws1) or 16
    _set_xlsx_counts(ws1, row, groups)
    # Дата в шапке
    ws1["B1"].value = f"Строевая записка \nНогинского спасательного центра на {date_s} "

    ws2 = wb["Лист2"]
    row2 = _find_1125_row(ws2) or 9
    _write_lист2_komandirovka(ws2, row2, groups.get("командировка", []))

    wb.save(dst)
    wb.close()


def fill_narrow_xlsx(dst: Path, people: list[dict], date_s: str) -> None:
    shutil.copy2(TEMPLATE_NARROW_XLSX, dst)
    wb = openpyxl.load_workbook(dst)
    groups = group_by_status(people)

    ws1 = wb["Лист1"]
    row = _find_1125_row(ws1) or 9
    _set_xlsx_counts(ws1, row, groups)
    ws1["B1"].value = f"Строевая записка \nНогинского спасательного центра на {date_s} "

    ws2 = wb["Лист2"]
    row2 = _find_1125_row(ws2) or 9
    _write_lист2_komandirovka(ws2, row2, groups.get("командировка", []))

    wb.save(dst)
    wb.close()


def fill_shds_doc(dst: Path, people: list[dict]) -> None:
    """6 человек из списка в блок «Группа РХБЗ» (Word COM)."""
    import win32com.client

    # Берём первых 6 из общего списка (или всех работающих, если хватает)
    working = [p for p in people if not p.get("status")]
    chosen = (working if len(working) >= 6 else people)[:6]

    shutil.copy2(TEMPLATE_SHDS_DOC, dst)
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(str(dst.resolve()))
    table = doc.Tables(1)

    for i, row_i in enumerate(SHDS_RHBZ_ROWS):
        if i >= len(chosen):
            break
        p = chosen[i]
        pos = SHDS_DEFAULT_POSITIONS[i]
        try:
            table.Cell(row_i, 2).Range.Text = pos
            table.Cell(row_i, 3).Range.Text = p.get("fio", "")
            table.Cell(row_i, 5).Range.Text = "1125 ЦРХБЗ"
        except Exception:
            pass

    doc.Save()
    doc.Close(False)
    word.Quit()


def generate_all() -> dict:
    """
    Создать 4 файла в {work_dir}/{дата}/.
    Возвращает пути и статистику.
    """
    ok, msg = storage.require_work_dir()
    if not ok:
        return {"ok": False, "error": msg}

    people = storage.load_people()
    if not people:
        return {"ok": False, "error": "Список людей пуст."}

    out = output_dir()
    date_s = today_str()
    files = {}

    dst_docx = out / f"Строевая записка {date_s}.docx"
    fill_stroevaya_docx(dst_docx, people, date_s)
    files["stroevaya_docx"] = str(dst_docx)

    dst_full = out / f"Строевая записка {date_s}.xlsx"
    fill_full_xlsx(dst_full, people, date_s)
    files["stroevaya_full_xlsx"] = str(dst_full)

    dst_narrow = out / f"Строевая записка {date_s} 1125 Ц(РХБЗ).xlsx"
    fill_narrow_xlsx(dst_narrow, people, date_s)
    files["stroevaya_narrow_xlsx"] = str(dst_narrow)

    dst_shds = out / f"ШДС СО 30 чел (ОРФОШ) {date_s}.doc"
    fill_shds_doc(dst_shds, people)
    files["shds_doc"] = str(dst_shds)

    return {
        "ok": True,
        "folder": str(out),
        "date": date_s,
        "people": len(people),
        "files": files,
    }
